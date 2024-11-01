# utils/transcript_processor.py
import io
import webvtt
from docx import Document
from typing import Tuple, Optional
import streamlit as st

class TranscriptProcessor:
    @staticmethod
    def process_docx(file_data: bytes) -> Optional[str]:
        """Process DOCX file and extract text content"""
        try:
            doc = Document(io.BytesIO(file_data))
            # Extract text from paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only include non-empty paragraphs
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            st.error(f"Error processing DOCX file: {str(e)}")
            return None

    @staticmethod
    def process_vtt(file_data: bytes) -> Optional[str]:
        """Process VTT file and extract text content"""
        try:
            # Write bytes to a temporary file since webvtt-py requires a file
            temp_file = io.StringIO(file_data.decode('utf-8'))
            
            # Parse VTT content
            full_text = []
            current_speaker = None
            
            for line in temp_file.getvalue().split('\n'):
                line = line.strip()
                
                # Skip empty lines, timestamps, and VTT header
                if not line or line == 'WEBVTT' or '-->' in line or line.isdigit():
                    continue
                
                # Check for speaker change (usually in format "SPEAKER:")
                if line.endswith(':') and len(line.split()) <= 2:
                    current_speaker = line
                    continue
                
                # Add speaker label if available
                if current_speaker and line:
                    full_text.append(f"{current_speaker} {line}")
                elif line:
                    full_text.append(line)
            
            return '\n'.join(full_text)
        except Exception as e:
            st.error(f"Error processing VTT file: {str(e)}")
            return None

    @staticmethod
    def process_txt(file_data: bytes) -> Optional[str]:
        """Process TXT file and extract content"""
        try:
            return file_data.decode('utf-8')
        except Exception as e:
            st.error(f"Error processing TXT file: {str(e)}")
            return None

    @classmethod
    def process_transcript_file(cls, uploaded_file) -> Tuple[str, bool]:
        """
        Process uploaded transcript file and return its content
        
        Args:
            uploaded_file: Streamlit UploadedFile object
        
        Returns:
            Tuple of (file_content: str, success: bool)
        """
        try:
            # Get file extension
            file_extension = uploaded_file.name.split('.')[-1].lower()
            file_data = uploaded_file.read()
            
            # Process based on file type
            if file_extension == 'docx':
                content = cls.process_docx(file_data)
            elif file_extension == 'vtt':
                content = cls.process_vtt(file_data)
            elif file_extension == 'txt':
                content = cls.process_txt(file_data)
            else:
                st.error(f"Unsupported file format: {file_extension}")
                return "", False
            
            if content:
                return content, True
            return "", False
            
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            return "", False

    @staticmethod
    def clean_transcript(text: str) -> str:
        """
        Clean and normalize transcript text
        
        Args:
            text: Raw transcript text
            
        Returns:
            Cleaned and normalized text
        """
        if not text:
            return text
            
        # Split into lines and remove empty ones
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Remove common transcript artifacts
        cleaned_lines = []
        for line in lines:
            # Remove timestamp patterns (various formats)
            if not any(pattern in line.lower() for pattern in ['-->', '[0-9]:', '(silence)']):
                cleaned_lines.append(line)
        
        # Join lines back together
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove multiple spaces
        cleaned_text = ' '.join(cleaned_text.split())
        
        return cleaned_text